from django.db import models
from django.db.models.signals import post_save
from django.dispatch import receiver
import pandas as pd
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.core.files.storage import FileSystemStorage
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from datetime import date
from django.core.files import File
import os
import shutil

fs = FileSystemStorage(location='/media')                   # to set the default storage path

# these classes are file managers and the variables are attributes which are stored in database


class Dsaform(models.Model):
    excel_file = models.FileField(storage=fs)
    excel_text = models.CharField(max_length=100, default='Enter Department Name')

    def __str__(self):
        return '%d %s ' % (self.id, self.excel_text)


class Dsauserform(models.Model):
    teacher_name = models.CharField(max_length=100, default='')


class Dsauser(models.Model):
    userfile_name = models.CharField(max_length=100, default='')
    user_file = models.FileField(upload_to='uploaded_docx_files/', default='')

    def __str__(self):
        return '%s %d' % (self.userfile_name, self.id)


def eng_to_nep():
    a = date.today()
    y = int(a.year)
    m = int(a.month)
    d = int(a.day)

    d = 16 + d

    if d > 30:
        m = m + 9
        d = d % 30
    else:
        m = m + 8

    if m > 12:
        y = y + 57
        m = m % 12
    else:
        y = y + 56

    return str(y) + '-' + str(m) + "-" + str(d)

# receiver tag brings the execution here no matter where it is


@receiver(post_save, sender=Dsaform)
def read_excel_file(sender, instance, created, *args, **kwargs):
    print(instance.excel_file.path)
    df = pd.read_excel(instance.excel_file.path)
    pd.set_option('display.expand_frame_repr', False)
    xls = pd.ExcelFile(instance.excel_file.path)
    xlsheet = xls.sheet_names
    xlslen = xlsheet.__len__()
    sort_excel(xls, xlslen, xlsheet)


def sort_excel(df, xlslen, xlsheet):                   # to sort the read xlsfile to datastructures

    for i in range(0, xlslen):

        dfs = pd.read_excel(df, sheet_name=xlsheet[i])
        dfs = dfs.fillna('')
        # df1 = dfsa.iloc[:, 7:13]

        sem_name = str(dfs.iloc[3][0])                  # to find the semester is odd or even
        if sem_name.find('-01-'):
            sem_name = 'EVEN'
        else:
            sem_name = 'ODD'
        program_name = str(dfs.iloc[3][0])
        program_name = program_name.replace('Program = ', '')
        program_name = program_name.replace(',', ' & ')

        period_total = str(dfs.iloc[3][3])                      # to find the total number of periods in a week
        period_total = period_total.replace('Total Classes Per Week =', '')
        period_total = period_total.rpartition('=')[2]          # breaks string into 2parts and stores the latter one

        teacher_name = str(dfs.iloc[4][0])
        teacher_name = teacher_name.replace('Teacher =', '')
# keeps only the teachers name from the string

        dep_name = str(dfs.iloc[1][0])
        dep_name = dep_name.replace('Department of', '')         # keeps only the department name from the string

        sub_names = ['']
        for j in range(7, 13):                                  # to get all subject names
            for k in range(1, 9):
                name = dfs.iloc[j][k]
                sub_names.append(str(name))

        sub_name = ['']
        for i in range(0, sub_names.__len__()):
            subname = str(sub_names[i])
            sub_name.append(subname)

        for elem in sub_name:                                # removes Break since it is not counted as subject
            if elem == 'Break':
                sub_name.remove(elem)

        while '' in sub_name:
            sub_name.remove('')

        for i in range(0, sub_name.__len__()):
            sub_name[i] = sub_name[i].replace('\n', '')

        for i in range(0, sub_name.__len__()):                                    # checks if subject is elective
            if '[ELECTIVEI]' in sub_name[i] or '[ELECTIVE I]' in sub_name[i]:
                sub_name[i] = sub_name[i].replace('[ELECTIVEI]', '')
                sub_name[i] = sub_name[i].replace('[ELECTIVE I]', '')
                sub_name[i] = 'Elective-III ' + sub_name[i]
            elif '[ELECTIVEII]' in sub_name[i] or '[ELECTIVE II]' in sub_name[i]:
                sub_name[i] = sub_name[i].replace('[ELECTIVEII]', '')
                sub_name[i] = sub_name[i].replace('[ELECTIVE II]', '')
                sub_name[i] = 'Elective-III ' + sub_name[i]
            elif '[ELECTIVEIII]' in sub_name[i] or '[ELECTIVE III]' in sub_name[i]:
                sub_name[i] = sub_name[i].replace('[ELECTIVEIII]', '')
                sub_name[i] = sub_name[i].replace('[ELECTIVE III]', '')
                sub_name[i] = 'Elective-III ' + sub_name[i]
            elif 'Elective' in sub_name:
                pas = sub_name[i]
                pos = pas.find('(')
                pas[pos] = ''
                pos = pas.find(')')
                pas[pos] = ''
                sub_name[i] = pas

            else:
                sub_name[i] = sub_name[i]

        for i in range(0, sub_name.__len__()):
            pan = sub_name[i]
            if 'BCT' in pan:
                pin = pan.find('[')
                sub_name[i] = pan[:pin] + '[BCT]' + pan[pin:]
            elif 'BEX' in pan:
                pin = pan.find('[')
                sub_name[i] = pan[:pin] + '[BEX]' + pan[pin:]
            elif 'BEL' in pan:
                pin = pan.find('[')
                sub_name[i] = pan[:pin] + '[BEL]' + pan[pin:]
            elif 'BCE' in pan:
                pin = pan.find('[')
                sub_name[i] = pan[:pin] + '[BCE]' + pan[pin:]
            elif 'BME' in pan:
                pin = pan.find('[')
                sub_name[i] = pan[:pin] + '[BME]' + pan[pin:]
            elif 'B.Arch' in pan:
                pin = pan.find('[')
                sub_name[i] = pan[:pin] + '[B.Arch]' + pan[pin:]
            else:
                sub_name[i] = sub_name[i]
        for i in range(0, sub_name.__len__()):
            sub_name[i] = sub_name[i].replace('[Lecture]', '[L]')
            sub_name[i] = sub_name[i].replace('[Tutorial]', '[T]')
            sub_name[i] = sub_name[i].replace('[Practical]', '[P]')
            sub_name[i] = sub_name[i].replace('[Lecture + Tutorial]', '[L+T]')
            if 'Project' in sub_name[i] and 'LAB' in sub_name[i]:
                pan = sub_name[i]
                pin = pan.find('(')
                sub_name[i] = pan[:pin] + '[P]' + pan[pin:]
            elif 'Project' in sub_name[i] and 'LAB' in sub_name[i]:
                pan = sub_name[i]
                pin = pan.find('(')
                sub_name[i] = pan[:pin] + '[P]' + pan[pin:]
            elif 'Project' in sub_name[i] and 'LAB' in sub_name[i]:
                pan = sub_name[i]
                pin = pan.find('(')
                sub_name[i] = pan[:pin] + '[P]' + pan[pin:]
            else:
                sub_name[i] = sub_name[i]

            sep = '('
            sub_name[i] = sub_name[i].split(sep, 1)[0]                      # the ultimate separation


# if a subject is lecture+tutorial, separate them and count them individually
        for i in range(0, sub_name.__len__()):
            if '(ALTERNATE WEEK)' in sub_name[i]:
                sub_name[i] = sub_name[i].replace('(ALTERNATE WEEK)', 'ALTERNATE WEEK')

            if '[L+T]' in sub_name[i]:
                x = sub_name[i].replace('[L+T]', '[L]')
                sub_name.append(x)
                x = sub_name[i].replace('[L+T]', '[T]')
                sub_name.append(x)
                sub_name.remove(sub_name[i])

        temp1 = []
        temp2 = []
        temp3 = []
        temp4 = []
        sub_names = []
        subnames_count = []
        subname_count = [[x, sub_name.count(x)] for x in set(sub_name)]

        for elem in subname_count:
            if '[L]' in elem[0]:
                temp1.append(elem)
            elif '[T]' in elem[0]:
                temp2.append(elem)
            elif '[P]' in elem[0]:
                temp3.append(elem)
            else:
                temp4.append(elem)
        subname_count = temp1 + temp2 + temp3 + temp4

        for i in range(0, subname_count.__len__()):
            item = subname_count[i]
            sub_names.append(item[0])
            subnames_count.append(str(item[1]))

        sub_name = sub_names
        no_subs = sub_name.__len__()

# this part exists to count the number of days a teacher has classes
        class_no = 0
        for i in range(7, 13):
            class_day = list(dfs.iloc[i])
            for elem in class_day:
                if elem == 'Sunday' or elem == 'Monday' or elem == 'Tuesday' or elem == 'Wednesday' or \
                        elem == 'Thursday' or elem == 'Friday':
                    class_day.remove(elem)  # to delete the day name from the rad string
            if class_day:
                class_no = class_no + 1

        create_doc(str(class_no), sem_name, teacher_name, dep_name, str(no_subs), period_total, sub_name, subnames_count)

# this function creates the doc file from the variables


def create_doc(class_no, sem_name, teacher_name, dep_name, no_subs, period_total, sub_name, subnames_count):

    no_subs = int(no_subs)
    if 'Prof. Dr.' in teacher_name:
        post = 'Professor'
    else:
        post = 'Teacher'

    records = ((),)
# to set the class,teacher,period and student numbers for lecture,tutorial and practical
    period_num = 1
    for i in range(0, no_subs):
        if '[T]' in sub_name[i]:
            class_type = 1
            teacher_num = 1
            period_num = 2
            student_num = 24

        elif '[L]' in sub_name[i]:
            class_type = 1
            teacher_num = 1
            period_num = 3
            student_num = 48
        elif '[P]' in sub_name[i] and ('Drawing' in sub_name[i] or 'Design' in sub_name[i]
                                       or 'Studio' in sub_name[i] or 'Paper work' in sub_name[i]):
            class_type = 2
            teacher_num = 3
            student_num = 24
            period_num = 3 * int(subnames_count[i])

        elif 'Project' in sub_name[i]:
            period_num = 3 * int(subnames_count[i])
            class_type = 1
            teacher_num = 1
            student_num = 4

        elif 'ALTERNATE WEEK' in sub_name[i]:
            class_type = 3
            period_num = 0.5
            teacher_num = 1
            student_num = 24
            sub_name[i] = sub_name[i].replace('ALTERNATE WEEK','')
        else:
            class_type = 3
            teacher_num = 3
            student_num = 24
            period_num = 3

        sn = str(i+1)
        result = (sn, str(sub_name[i]), str(class_type), str(teacher_num), str(period_num), str(student_num))
        records = records + (result,)                   # adds the tuple result to the tuple of tuples i.e.  records


# divide records to two parts and keep latter;
    records = records[1:]                 # partition from second element because defining stored empty tuple at first




    document = Document()

    #Different styles
    obj_charstyle = document.styles.add_style('nepali', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(12)
    obj_font.name = 'Preeti'

    obj_charstyle = document.styles.add_style('parnepali', WD_STYLE_TYPE.PARAGRAPH)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(12)
    obj_font.name = 'Preeti'

    obj_charstyle = document.styles.add_style('english', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(11)
    obj_font.name = 'Calibri Body'

    obj_font = document.styles['Normal'].font
    obj_font.name='Calibri Body'

    #header at top
    paragraph=document.add_paragraph()
    paragraph.style='parnepali'
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run('lqe\'jg ljZjljBfno\n').bold=True
    paragraph.add_run('OlGhlgol/Gª cWoog ;+:yfg\n').bold=True
    paragraph.add_run('s]lGb|o SofDk; k\'Nrf]s\n').bold=True
    paragraph.add_run('lzIfs sfo{ ;Dkfbg kmf/fd').bold=True

    #Introduction table

    table1=['ljefu ','Semester ','lzIfssf] gfd y/ ','kb ','k|lt xKtf sIff ePsf] lbg ']
    table2=[dep_name, sem_name, teacher_name, post, str(class_no)]
    table=document.add_table(3,2,style = 'Table Grid')
    table.autofit=True
    table.cell(0,0).width=Inches(6)
    table.cell(0,1).width=Inches(2)
    table.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    table.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for i in range(int(len(table1))):
        row=table.rows[int(i/2)]
        paragraph=row.cells[i%2].paragraphs[0]
        if i%2==1:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if i==1 :
            paragraph.add_run(table1[i]+" : ").bold=True
        else:
            paragraph.add_run(table1[i]+" M ",style='nepali').bold=True
        paragraph.add_run(table2[i]).bold=True
    paragraph=document.add_paragraph()
    paragraph.add_run("\n:gfts tx sIff ljj/0f",style='nepali').underline=True

    #our main table

    table1=['qm=;+=','ljifo','sIff lsl;d',';+nUg lzIfs ;+Vof','lkl/o8','ljBfyL{ ;+Vof']
    table = document.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    table.autofit =True

    table.cell(0,0).width=Inches(1)
    table.cell(0,1).width=Inches(14)
    table.cell(0,2).width=Inches(2)
    table.cell(0,3).width=Inches(5)
    table.cell(0,4).width=Inches(7)
    table.cell(0,5).width=Inches(1)

    for i in range(len(table1)):
        cell=table.cell(0,i)
        paragraph=cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(table1[i],style='nepali').bold=True

    total = 0

    for snid, subid, classid, teacherid, periodid, studentid in records:
        row_cells = table.add_row().cells
        for i in range(5):
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        s=row_cells[0].paragraphs[0].add_run(snid)
        s.bold=True
        s.font.size=Pt(10)
        s=row_cells[1].paragraphs[0].add_run(subid)
        s.bold=True
        s.font.size=Pt(10)
        s=row_cells[2].paragraphs[0].add_run(classid)
        s.bold=True
        s.font.size=Pt(10)
        s=row_cells[3].paragraphs[0].add_run(teacherid)
        s.bold=True
        s.font.size=Pt(10)
        s=row_cells[4].paragraphs[0].add_run(periodid)
        s.bold=True
        s.font.size=Pt(10)
        s=row_cells[5].paragraphs[0].add_run(studentid)
        s.bold=True
        s.font.size=Pt(10)
        total=total + int(periodid)
    cell=table.add_row().cells
    paragraph="Total = "+str(total)+" Periods"
    cell[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    s=cell[4].paragraphs[0].add_run(paragraph)
    s.bold=True
    s.font.size=Pt(9)
    paragraph=document.add_paragraph()
    paragraph.add_run("\n:gftsf]Q/ tx sIff ljj/0f",style='nepali').underline=True

    #Master level Table
    table1=['qm=;+=','ljifo',';+nUg lzIfs ;+Vof','q]ml86 cfj/']
    table = document.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.autofit = True
    table.cell(0,1).width=Inches(4)
    for i in range(len(table1)):
        cell=table.cell(0,i).paragraphs[0]
        cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.add_run(table1[i],style='nepali').bold=True
    table.add_row()

    #Notes and periods
    s = document.add_paragraph()
    s.add_run("\nGff]6 M    -s_ lkl/o8 M kf7|oqmddf pNn]v eP adf]lhd x'g]5 .\n",style='nepali')
    s.add_run(" \t-v_ sIff lsl;d eGgfn]",style='nepali')
    s.add_run(" 1,2,3 ").bold=True
    s.add_run("hgfpg' kg]{5 .\n",style='nepali')
    s.add_run("\t   1)").bold=True
    s.add_run("eGgfn]",style='nepali')
    s.add_run(" Theory /Tutorial /B.E. Project /B.Arch. Thesis\n\t   2)").bold=True
    s.add_run(" eGgfn]",style='nepali')
    s.add_run(" Drawing /Design /Design Studio /Paper work").bold=True
    s.add_run(" x'g] Nofj\n",style='nepali')
    s.add_run("\t   3)").bold=True
    s.add_run(" eGgfn] ",style='nepali')
    s.add_run("2").bold=True
    s.add_run(" df pNn]v gePsf Nofjx?",style='nepali')

    document.add_paragraph()
    #document.add_page_break()

    table1=['lzIfssf] x:tfIf/ M =========================================','k|dfl0ft ug]{','lzIfssf] gfd y/ M ','laefuLo k|d\'v','ldtL M ']
    #table2=["","", teacher.name, "", str(date.today())]
    table2=["","", teacher_name, "", eng_to_nep()]



    table=document.add_table(3,2)
    table.autofit=True
    table.cell(0,0).width=Inches(10)
    table.cell(0,1).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
    table.cell(1,1).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
    for i in range(len(table1)):
        cell=table.cell(int(i/2),i%2)
        paragraph=cell.paragraphs[0]
        paragraph.add_run(table1[i],style='nepali')
        paragraph.add_run(table2[i]).bold=True

    #footnote
    paragraph = document.add_paragraph()
    paragraph.page_break_before = True
    paragraph.style='parnepali'
    paragraph.add_run("\n\n\n\nb|i6Jo M\t!_ sIff ?l6g ;+nUg x'g''kg]{5 .	 @_ ")
    paragraph.add_run("Elective Course ",style='english').bold=True
    paragraph.add_run("sf nflu ljBfyL{ ;+Vof $* x'g]5 .\n\t#_")
    paragraph.add_run(" Master/Ph.D. ",style='english').bold=True
    paragraph.add_run("sf]")
    paragraph.add_run(" Thesis ",style='english').bold=True
    paragraph.add_run("sf nflu of] kmd{ eg{ cfjZos 5}g .")

    document.add_page_break()
    # saves teacher name in database in dsauser model
    # saves document with the teacher name as file name
    teacher_name = teacher_name.replace(' ', '')
    document.save('media/files/' + str(teacher_name.lower()) + '.docx')
    f = open('media/files/' + str(teacher_name.lower()) + '.docx', 'rb')

    if Dsauser.objects.filter(userfile_name=teacher_name.lower()).exists():     # to replace files if they exist already
        dsauser = Dsauser.objects.filter(userfile_name=teacher_name.lower())
        dsauser_obj = dsauser.first()
        file = dsauser_obj.user_file.path
        dsauser.user_file = File(f)
        os.remove(file)
    else:
        dsa = Dsauser(userfile_name=teacher_name.lower(), user_file=File(f))
        dsa.save()




